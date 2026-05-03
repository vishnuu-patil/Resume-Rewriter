from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import requests
import base64
import os
from datetime import datetime
from dotenv import load_dotenv
load_dotenv()

app = Flask(__name__)

TEAL = RGBColor(0x00, 0x7B, 0x83)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '007B83')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_font(run, size, bold=False, color=DARK, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    run.font.name = 'Calibri'

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    set_font(run, 11, bold=True, color=TEAL)
    add_horizontal_line(p)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def build_docx(data, output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data.get('name', '').upper())
    set_font(name_run, 20, bold=True, color=TEAL)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(data.get('title', ''))
    set_font(title_run, 11, italic=True)

    # Contact
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run(data.get('contact', ''))
    set_font(contact_run, 10, color=TEAL)

    # Summary
    if data.get('summary'):
        add_section_heading(doc, 'Professional Summary')
        p = doc.add_paragraph()
        run = p.add_run(data['summary'])
        set_font(run, 10)

    # Core Competencies
    if data.get('competencies'):
        add_section_heading(doc, 'Core Competencies')
        items = data['competencies']
        rows = (len(items) + 2) // 3
        table = doc.add_table(rows=rows, cols=3)
        table.style = 'Table Grid'
        idx = 0
        for row in table.rows:
            for cell in row.cells:
                if idx < len(items):
                    cell.text = items[idx]
                    for para in cell.paragraphs:
                        for run in para.runs:
                            set_font(run, 10, bold=True, color=TEAL)
                    idx += 1

    # Technical Skills
    if data.get('skills'):
        add_section_heading(doc, 'Technical Skills')
        skills = data['skills']
        table = doc.add_table(rows=len(skills), cols=2)
        table.style = 'Table Grid'
        for i, skill in enumerate(skills):
            row = table.rows[i]
            row.cells[0].text = skill.get('category', '')
            row.cells[1].text = skill.get('detail', '')
            for run in row.cells[0].paragraphs[0].runs:
                set_font(run, 10, bold=True, color=TEAL)
            for run in row.cells[1].paragraphs[0].runs:
                set_font(run, 10)

    # Work Experience
    if data.get('experience'):
        add_section_heading(doc, 'Work Experience')
        for exp in data['experience']:
            p = doc.add_paragraph()
            role_run = p.add_run(exp.get('role', ''))
            set_font(role_run, 10, bold=True)
            company_run = p.add_run(f" | {exp.get('company', '')}")
            set_font(company_run, 10, italic=True)
            date_run = p.add_run(f"  {exp.get('dates', '')}")
            set_font(date_run, 10, italic=True)
            for bullet in exp.get('bullets', []):
                add_bullet(doc, bullet)

    # Projects
    if data.get('projects'):
        add_section_heading(doc, 'Key Projects')
        for proj in data['projects']:
            p = doc.add_paragraph()
            title_run = p.add_run(proj.get('title', ''))
            set_font(title_run, 10, bold=True)
            if proj.get('date'):
                date_run = p.add_run(f"  {proj['date']}")
                set_font(date_run, 10, italic=True)
            for bullet in proj.get('bullets', []):
                add_bullet(doc, bullet)

    # Certifications
    if data.get('certifications'):
        add_section_heading(doc, 'Certifications')
        for cert in data['certifications']:
            add_bullet(doc, cert)

    # Education
    if data.get('education'):
        add_section_heading(doc, 'Education')
        for edu in data['education']:
            p = doc.add_paragraph()
            deg_run = p.add_run(edu.get('degree', ''))
            set_font(deg_run, 10, bold=True)
            school_run = p.add_run(f" | {edu.get('school', '')}")
            set_font(school_run, 10, italic=True)
            date_run = p.add_run(f"  {edu.get('dates', '')}")
            set_font(date_run, 10, italic=True)

    # Achievements
    if data.get('achievements'):
        add_section_heading(doc, 'Achievements & Leadership')
        for ach in data['achievements']:
            add_bullet(doc, ach)

    doc.save(output_path)

def call_groq(prompt, system):
    response = requests.post(
        'https://api.groq.com/openai/v1/chat/completions',
        headers={
            'Authorization': f'Bearer {GROQ_API_KEY}',
            'Content-Type': 'application/json'
        },
        json={
            'model': 'llama-3.3-70b-versatile',
            'messages': [
                {'role': 'system', 'content': system},
                {'role': 'user', 'content': prompt}
            ],
            'temperature': 0.3
        }
    )
    return response.json()['choices'][0]['message']['content']

@app.route('/process', methods=['POST'])
def process_resume():
    data = request.json
    rewritten_text = data.get('rewritten_text', '')
    sender = data.get('sender', '')
    jd_text = data.get('jd_text', '')
    resume_text = data.get('resume_text', '')

    # Step 1 — Structure rewritten resume into JSON
    print("Structuring resume into JSON...")
    structure_prompt = f"""
Convert this rewritten resume into a JSON object with exactly this structure.
Return ONLY the JSON, no explanation, no markdown backticks.

{{
  "name": "Full Name",
  "title": "Job Title | Other Title",
  "contact": "email • phone • city • LinkedIn • Portfolio",
  "summary": "paragraph text",
  "competencies": ["item1", "item2", "item3"],
  "skills": [
    {{"category": "Analytics", "detail": "Excel, SQL..."}},
    {{"category": "Database", "detail": "SQL..."}}
  ],
  "experience": [
    {{
      "role": "Job Title",
      "company": "Company Name, City",
      "dates": "Month Year – Present",
      "bullets": ["bullet 1", "bullet 2"]
    }}
  ],
  "projects": [
    {{
      "title": "Project Name | details",
      "date": "Month Year",
      "bullets": ["bullet 1", "bullet 2"]
    }}
  ],
  "certifications": ["cert 1", "cert 2"],
  "education": [
    {{
      "degree": "Degree Name",
      "school": "College Name",
      "dates": "Year – Year"
    }}
  ],
  "achievements": ["achievement 1", "achievement 2"]
}}

RESUME TEXT:
{rewritten_text}
"""
    structured_json = call_groq(structure_prompt, "You are a JSON formatter. Return only valid JSON.")

    try:
        resume_data = json.loads(structured_json)
    except:
        # Clean and retry
        cleaned = structured_json.strip().replace('```json', '').replace('```', '')
        resume_data = json.loads(cleaned)

    # Step 2 — Generate filename
    name = resume_data.get('name', 'Full Name')
    parts = name.strip().split()
    if len(parts) >= 2:
        firstname = parts[-1]
        lastname = parts[0]
    else:
        firstname = parts[0]
        lastname = 'User'
    month = datetime.now().strftime('%b')
    year = datetime.now().strftime('%Y')
    filename = f"{lastname}_{firstname}_{month}_{year}.docx"
    output_path = os.path.join('outputs', filename)
    os.makedirs('outputs', exist_ok=True)

    # Step 3 — Build DOCX
    print(f"Building DOCX: {filename}")
    build_docx(resume_data, output_path)

    # Step 4 — Gap analysis
    print("Running gap analysis...")
    gap_prompt = f"""
Analyze this resume against the job description and return a JSON with exactly this structure.
Return ONLY JSON, no explanation.

{{
  "match_score": 75,
  "matched_keywords": ["keyword1", "keyword2"],
  "missing_skills": ["skill1", "skill2"],
  "missing_keywords": ["keyword1", "keyword2"],
  "tips": ["tip1", "tip2", "tip3"]
}}

JOB DESCRIPTION:
{jd_text}

RESUME:
{resume_text}
"""
    gap_json = call_groq(gap_prompt, "You are a resume analyst. Return only valid JSON.")

    try:
        gap_data = json.loads(gap_json)
    except:
        cleaned = gap_json.strip().replace('```json', '').replace('```', '')
        gap_data = json.loads(cleaned)

    # Step 5 — Read DOCX as base64 to send back
    with open(output_path, 'rb') as f:
        docx_base64 = base64.b64encode(f.read()).decode('utf-8')

    return jsonify({
        'filename': filename,
        'docx_base64': docx_base64,
        'match_score': gap_data.get('match_score', 0),
        'matched_keywords': gap_data.get('matched_keywords', []),
        'missing_skills': gap_data.get('missing_skills', []),
        'missing_keywords': gap_data.get('missing_keywords', []),
        'tips': gap_data.get('tips', []),
        'sender': sender
    })

if __name__ == '__main__':
    app.run(port=5000, debug=True)